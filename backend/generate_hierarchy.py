from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()

# Add title
doc.add_heading('Expense Tracker App - Project Hierarchy', 0)

# Add architectural pattern explanation
doc.add_paragraph('This document outlines the project hierarchy based on Layered Architecture pattern.')
doc.add_paragraph('Layered Architecture divides the application into layers: Presentation, Application, and Data.')

# Presentation Layer
doc.add_heading('Presentation Layer (Frontend)', level=1)
doc.add_paragraph('This layer handles user interface and interactions.')
doc.add_paragraph('Location: src/')
doc.add_paragraph('- App.jsx: Main application component')
doc.add_paragraph('- main.jsx: Entry point')
doc.add_paragraph('- index.css: Global styles')
doc.add_paragraph('- App.css: Application styles')
doc.add_paragraph('- components/:')
doc.add_paragraph('  - Dashboard/: Dashboard component and styles')
doc.add_paragraph('  - Expense/: Expense management component and styles')
doc.add_paragraph('  - Income/: Income management component and styles')
doc.add_paragraph('  - SignIn/: User sign-in component and styles')
doc.add_paragraph('  - SignUp/: User sign-up component and styles')
doc.add_paragraph('  - ForgotPassword/: Password recovery components and styles')
doc.add_paragraph('  - Other components: DataTable, Form, InputField, Layout, Navbar, Routes')

# Application Layer
doc.add_heading('Application Layer (Backend)', level=1)
doc.add_paragraph('This layer contains business logic and API endpoints.')
doc.add_paragraph('Location: backend/')
doc.add_paragraph('- main.py: FastAPI application and API routes')
doc.add_paragraph('- models.py: Data models (SQLAlchemy)')
doc.add_paragraph('- database.py: Database connection and session management')
doc.add_paragraph('- requirements.txt: Python dependencies')

# Data Layer
doc.add_heading('Data Layer', level=1)
doc.add_paragraph('This layer manages data storage and retrieval.')
doc.add_paragraph('- expense_tracker.db: SQLite database file')
doc.add_paragraph('- Database schema defined in models.py')

# Additional files
doc.add_heading('Additional Configuration Files', level=1)
doc.add_paragraph('- package.json: Frontend dependencies and scripts')
doc.add_paragraph('- vite.config.js: Vite configuration')
doc.add_paragraph('- eslint.config.js: ESLint configuration')
doc.add_paragraph('- .gitignore: Git ignore rules')
doc.add_paragraph('- README.md: Project documentation')
doc.add_paragraph('- TECHNICAL_DOCUMENTATION.txt: Technical details')
doc.add_paragraph('- Other files: .env, validations.txt, etc.')

# Save the document
doc.save('project_hierarchy.docx')

print("Word document 'project_hierarchy.docx' created successfully.")
